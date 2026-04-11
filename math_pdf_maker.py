import sys
import os
import re
import json
import urllib.request
import html
import subprocess
import uuid
from datetime import datetime

# --- PRE-FLIGHT DEPENDENCY CHECK ---
def check_dependencies():
    missing = []
    try:
        import PySide6
    except ImportError:
        missing.append("PySide6")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    try:
        import rapidocr_onnxruntime
    except ImportError:
        missing.append("rapidocr-onnxruntime")
    try:
        import fitz  # PyMuPDF
    except ImportError:
        missing.append("PyMuPDF")
    try:
        from google import genai
    except ImportError:
        missing.append("google-genai")
    try:
        from PIL import Image
    except ImportError:
        missing.append("Pillow")
    try:
        import win32com.client
    except ImportError:
        missing.append("pywin32")
    
    if missing:
        print(f"Installing missing dependencies: {', '.join(missing)}...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-U"] + missing)
            print("Dependencies installed. Restarting...")
            os.execv(sys.executable, [sys.executable] + sys.argv)
        except Exception as e:
            print(f"CRITICAL ERROR: Failed to install dependencies: {e}")
            sys.exit(1)

check_dependencies()

from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                               QHBoxLayout, QSplitter, QPlainTextEdit, QPushButton, 
                               QLabel, QSpinBox, QDoubleSpinBox, QComboBox, 
                               QFileDialog, QMessageBox, QFrame, QLineEdit, QProgressBar)
from PySide6.QtWebEngineWidgets import QWebEngineView
from PySide6.QtGui import QPageLayout, QPageSize, QIcon, QFont
from PySide6.QtCore import Qt, QTimer, QUrl, QMargins, QMarginsF

# --- OFFLINE MATHJAX SETUP ---
APP_DIR = os.path.join(os.path.expanduser("~"), ".alvin_math_pdf")
MATHJAX_FILE = os.path.join(APP_DIR, "tex-svg.js") # FIX: Use SVG engine for flawless radical/square-root rendering

def ensure_offline_math_engine():
    """Downloads a standalone MathJax engine on the first run for complete offline use."""
    if not os.path.exists(APP_DIR):
        os.makedirs(APP_DIR)
    
    if not os.path.exists(MATHJAX_FILE):
        print("First run detected. Downloading offline MathJax engine (~1.5MB)...")
        url = "https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js"
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(MATHJAX_FILE, 'wb') as out_file:
                out_file.write(response.read())
            print("Offline engine ready!")
        except Exception as e:
            print(f"Failed to download MathJax: {e}")
            print("The app will attempt to use the live CDN fallback.")

ensure_offline_math_engine()


class MarkdownImageEditor(QPlainTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)
        # Create an images folder inside your existing app directory
        self.image_dir = os.path.join(APP_DIR, "images")
        if not os.path.exists(self.image_dir):
            os.makedirs(self.image_dir)

    def insertFromMimeData(self, source):
        # 1. Handle pure image (e.g., Right-click -> Copy Image or Snipping Tool)
        if source.hasImage():
            image = source.imageData()
            filename = f"pasted_img_{uuid.uuid4().hex[:8]}.png"
            filepath = os.path.join(self.image_dir, filename)
            image.save(filepath, "PNG")
            html_filepath = filepath.replace('\\', '/')
            self.insertPlainText(f"\n![Pasted Image](file:///{html_filepath})\n")
            return

        # 2. Handle Mixed Web Content (Highlighting text + images together)
        if source.hasHtml():
            html_data = source.html()
            # Look for external image links hidden inside the copied HTML code
            img_urls = re.findall(r'<img[^>]+src=["\'](https?://[^"\']+)["\']', html_data)
            
            if img_urls:
                # Paste the normal text first
                super().insertFromMimeData(source)
                self.insertPlainText("\n\n")
                
                # Automatically download the web images in the background
                for url in img_urls:
                    try:
                        filename = f"web_img_{uuid.uuid4().hex[:8]}.png"
                        filepath = os.path.join(self.image_dir, filename)
                        
                        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                        with urllib.request.urlopen(req, timeout=5) as response, open(filepath, 'wb') as f:
                            f.write(response.read())
                            
                        html_filepath = filepath.replace('\\', '/')
                        self.insertPlainText(f"![Downloaded Image](file:///{html_filepath})\n")
                    except Exception as e:
                        print(f"Background download failed: {e}")
                return

        # 3. Default fallback for standard plain text
        super().insertFromMimeData(source)


class MathPdfMaker(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Mr. Alvin's PDF Typesetter")
        self.setGeometry(100, 100, 1400, 800)
        self.setStyleSheet("QMainWindow { background-color: #1e1e1e; color: #ffffff; }")
        self.showMaximized() # FIX: Force window to scale and fit the screen properly on launch

        # Memory Cache Variables
        from PySide6.QtCore import QSettings
        self.settings = QSettings("MrAlvin", "MathPdfTypesetter")
        self.ocr_engine = None # Cache OCR in memory for instant subsequent scans

        # Variables
        self.mathjax_url = QUrl.fromLocalFile(MATHJAX_FILE).toString() if os.path.exists(MATHJAX_FILE) else "https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js"
        
        self.build_ui()
        
        # Debounce timer for live preview so typing doesn't stutter
        self.preview_timer = QTimer()
        self.preview_timer.setSingleShot(True)
        self.preview_timer.timeout.connect(self.update_preview)
        
        # Initial render
        self.update_preview()

    def build_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)

        # --- TOOLBAR (Top) ---
        toolbar = QFrame()
        toolbar.setStyleSheet("QFrame { background-color: #2e2e2e; border-radius: 5px; padding: 5px; } QLabel { color: #fff; font-weight: bold; }")
        tb_layout = QHBoxLayout(toolbar)
        
        # Header / Title
        tb_layout.addWidget(QLabel("Header:"))
        self.header_input = QLineEdit()
        self.header_input.setPlaceholderText("First Page Title...")
        self.header_input.setStyleSheet("background-color: #3e3e3e; color: #fff; padding: 2px;")
        self.header_input.textChanged.connect(lambda: self.preview_timer.start(600))
        tb_layout.addWidget(self.header_input)

        # Watermark
        tb_layout.addWidget(QLabel(" Watermark:"))
        self.watermark_input = QLineEdit()
        self.watermark_input.setPlaceholderText("Background text...")
        self.watermark_input.setStyleSheet("background-color: #3e3e3e; color: #fff; padding: 2px;")
        self.watermark_input.textChanged.connect(lambda: self.preview_timer.start(600))
        tb_layout.addWidget(self.watermark_input)

        # Page Size
        tb_layout.addWidget(QLabel("Page Size:"))
        self.page_size_cb = QComboBox()
        self.page_size_cb.addItems(["A4", "Letter", "Legal"])
        self.page_size_cb.currentTextChanged.connect(self.update_preview)
        tb_layout.addWidget(self.page_size_cb)

        # Margins
        tb_layout.addWidget(QLabel("  Margins (mm):"))
        self.margin_spin = QSpinBox()
        self.margin_spin.setRange(0, 50)
        self.margin_spin.setValue(20)
        self.margin_spin.valueChanged.connect(self.update_preview)
        tb_layout.addWidget(self.margin_spin)

        # Font Size
        tb_layout.addWidget(QLabel("  Font Size (pt):"))
        self.font_spin = QSpinBox()
        self.font_spin.setRange(8, 36)
        self.font_spin.setValue(11)
        self.font_spin.valueChanged.connect(self.update_preview)
        tb_layout.addWidget(self.font_spin)

        # Line Spacing
        tb_layout.addWidget(QLabel("  Line Spacing:"))
        self.line_space_spin = QDoubleSpinBox()
        self.line_space_spin.setRange(1.0, 3.0)
        self.line_space_spin.setSingleStep(0.1)
        self.line_space_spin.setValue(1.5)
        self.line_space_spin.valueChanged.connect(self.update_preview)
        tb_layout.addWidget(self.line_space_spin)
        
        tb_layout.addStretch()
        main_layout.addWidget(toolbar)

        # --- SPLITTER (Editor / Preview) ---
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter, stretch=1)

        # Editor Setup
        editor_widget = QWidget()
        editor_layout = QVBoxLayout(editor_widget)
        editor_layout.setContentsMargins(0, 0, 0, 0)
        
        lbl_layout = QHBoxLayout()
        lbl_layout.addWidget(QLabel("Raw Input (Markdown + LaTeX Math):"))
        lbl_layout.addStretch()
        
        # New Feature: Quick Math Snippets UI Toolbar
        btn_matrix = QPushButton("[::] bmatrix")
        btn_matrix.setCursor(Qt.PointingHandCursor)
        btn_matrix.setStyleSheet("padding: 2px 8px; background-color: #3e3e3e; border-radius: 3px; font-weight: bold;")
        btn_matrix.clicked.connect(lambda: self.editor.insertPlainText("$$ \\begin{bmatrix}\n  1 & 0 \\\\\n  0 & 1\n\\end{bmatrix} $$"))
        lbl_layout.addWidget(btn_matrix)
        
        btn_frac = QPushButton("a/b frac")
        btn_frac.setCursor(Qt.PointingHandCursor)
        btn_frac.setStyleSheet("padding: 2px 8px; background-color: #3e3e3e; border-radius: 3px; font-weight: bold;")
        btn_frac.clicked.connect(lambda: self.editor.insertPlainText("\\frac{numerator}{denominator}"))
        lbl_layout.addWidget(btn_frac)
        
        btn_sqrt = QPushButton("√ sqrt")
        btn_sqrt.setCursor(Qt.PointingHandCursor)
        btn_sqrt.setStyleSheet("padding: 2px 8px; background-color: #3e3e3e; border-radius: 3px; font-weight: bold;")
        btn_sqrt.clicked.connect(lambda: self.editor.insertPlainText("\\sqrt{x}"))
        lbl_layout.addWidget(btn_sqrt)

        btn_inf = QPushButton("∞ inf")
        btn_inf.setCursor(Qt.PointingHandCursor)
        btn_inf.setStyleSheet("padding: 2px 8px; background-color: #3e3e3e; border-radius: 3px; font-weight: bold;")
        btn_inf.clicked.connect(lambda: self.editor.insertPlainText("\\infty"))
        lbl_layout.addWidget(btn_inf)

        editor_layout.addLayout(lbl_layout)

        # --- FIX: Move OCR Tools to a second row to remove minimum width restrictions ---
        ocr_layout = QHBoxLayout()
        ocr_layout.addStretch() # Push OCR buttons to the right side neatly
        
        btn_ocr = QPushButton("📸 Local Offline OCR")
        btn_ocr.setCursor(Qt.PointingHandCursor)
        btn_ocr.setStyleSheet("padding: 4px 10px; background-color: #d35400; color: white; border-radius: 3px; font-weight: bold;")
        btn_ocr.clicked.connect(self.scan_image_ocr)
        ocr_layout.addWidget(btn_ocr)
        
        ocr_layout.addSpacing(10)
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("Paste Gemini API Key...")
        self.api_key_input.setEchoMode(QLineEdit.Password)
        self.api_key_input.setStyleSheet("background-color: #3e3e3e; color: #fff; padding: 2px; width: 140px;")
        ocr_layout.addWidget(self.api_key_input)
        
        btn_ai_ocr = QPushButton("✨ Gemini AI OCR")
        btn_ai_ocr.setCursor(Qt.PointingHandCursor)
        btn_ai_ocr.setStyleSheet("padding: 4px 10px; background-color: #8e44ad; color: white; border-radius: 3px; font-weight: bold;")
        btn_ai_ocr.clicked.connect(self.scan_image_gemini)
        ocr_layout.addWidget(btn_ai_ocr)

        editor_layout.addLayout(ocr_layout)
        
        self.editor = MarkdownImageEditor()
        self.editor.setFont(QFont("Consolas", 11))
        self.editor.setStyleSheet("background-color: #1e1e1e; color: #d4d4d4; border: 1px solid #333;")
        self.editor.textChanged.connect(lambda: self.preview_timer.start(600)) # 600ms debounce
        editor_layout.addWidget(self.editor)
        splitter.addWidget(editor_widget)

        # Preview Setup
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        preview_layout.setContentsMargins(0, 0, 0, 0)
        preview_layout.addWidget(QLabel("Live PDF Preview:"))
        
        self.web_view = QWebEngineView()
        self.web_view.setStyleSheet("background-color: #ffffff;") # Keep white for accurate paper rep
        preview_layout.addWidget(self.web_view)
        splitter.addWidget(preview_widget)
        
        # FIX: Make the splitter handle visible and draggable in dark mode
        splitter.setHandleWidth(8)
        splitter.setStyleSheet("QSplitter::handle { background-color: #444; border-radius: 3px; margin: 0px 2px; } QSplitter::handle:hover { background-color: #666; }")
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 1)
        splitter.setSizes([600, 800])

        # --- EXPORT BAR (Bottom) ---
        export_bar = QFrame()
        export_bar.setStyleSheet("QFrame { background-color: #2e2e2e; border-radius: 5px; padding: 5px; }")
        ex_layout = QHBoxLayout(export_bar)
        
        ex_layout.addWidget(QLabel("Filename:", styleSheet="color: #fff;"))
        self.filename_input = QLineEdit("Math_Document")
        self.filename_input.setStyleSheet("background-color: #3e3e3e; color: #fff; padding: 4px;")
        ex_layout.addWidget(self.filename_input)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.hide() # Hidden by default until needed
        ex_layout.addWidget(self.progress_bar)
        
        self.export_btn = QPushButton("💾 Export PDF")
        self.export_btn.setStyleSheet("QPushButton { background-color: #0078D4; color: white; font-weight: bold; padding: 6px 15px; border-radius: 3px; } QPushButton:hover { background-color: #106EBE; }")
        self.export_btn.clicked.connect(self.export_pdf)
        ex_layout.addWidget(self.export_btn)
        
        # Wire up the PDF finish signal permanently to avoid RuntimeWarnings
        self.web_view.page().pdfPrintingFinished.connect(self._on_pdf_finished)

        from PySide6.QtWidgets import QCheckBox
        self.auto_math_cb = QCheckBox("Auto-Convert Word Math (Win)")
        self.auto_math_cb.setStyleSheet("color: #fff; font-weight: bold;")
        self.auto_math_cb.setChecked(True)
        self.auto_math_cb.setToolTip("Uses MS Word to automatically compile equations to native formatting.")
        ex_layout.addWidget(self.auto_math_cb)

        self.export_word_btn = QPushButton("📝 Export Word (.docx)")
        self.export_word_btn.setStyleSheet("QPushButton { background-color: #2b579a; color: white; font-weight: bold; padding: 6px 15px; border-radius: 3px; } QPushButton:hover { background-color: #3e6db5; }")
        self.export_word_btn.clicked.connect(self.export_word)
        ex_layout.addWidget(self.export_word_btn)

        self.export_lo_btn = QPushButton("📝 Export LibreOffice")
        self.export_lo_btn.setStyleSheet("QPushButton { background-color: #18a303; color: white; font-weight: bold; padding: 6px 15px; border-radius: 3px; } QPushButton:hover { background-color: #21c408; }")
        self.export_lo_btn.clicked.connect(self.export_libreoffice)
        ex_layout.addWidget(self.export_lo_btn)
        
        main_layout.addWidget(export_bar)

    def generate_html(self):
        raw_text = self.editor.toPlainText()
        
        # --- FIX: Auto-correct AI copy-paste errors ---
        # When copying from AI chats, \[ and \] often get stripped to [ and ]
        raw_text = re.sub(r'^[ \t]*\[[ \t]*$', r'\\[', raw_text, flags=re.MULTILINE)
        raw_text = re.sub(r'^[ \t]*\][ \t]*$', r'\\]', raw_text, flags=re.MULTILINE)
        
        safe_text = html.escape(raw_text)
        
        # --- FIX: Protect Math blocks from HTML <br> injection ---
        # The previous naive replace('\n', '<br>') destroyed multiline LaTeX matrices!
        math_blocks = []
        def math_repl(match):
            # Auto-fix single backslashes at the end of matrix lines
            fixed_math = re.sub(r'(?<!\\)\\\s*\n', r'\\\\\n', match.group(0))
            math_blocks.append(fixed_math)
            return f"__MATH_BLOCK_{len(math_blocks)-1}__"

        # Extract properly formatted display and inline math safely into placeholders FIRST
        safe_text = re.sub(r'\$\$.*?\$\$', math_repl, safe_text, flags=re.DOTALL)
        safe_text = re.sub(r'(?<!\$)\$.*?\$(?!\$)', math_repl, safe_text)
        safe_text = re.sub(r'\\\[.*?\\\]', math_repl, safe_text, flags=re.DOTALL)
        safe_text = re.sub(r'\\\(.*?\\\)', math_repl, safe_text)

        # --- FIX: Autocorrect stray Parentheses to LaTeX Inline Math AFTER extracting good math ---
        # Converts (A \subseteq B) directly into a math block so it doesn't corrupt nested parentheses
        def autocorrect_repl(match):
            math_blocks.append(f"\\({match.group(1)}\\)")
            return f"__MATH_BLOCK_{len(math_blocks)-1}__"
            
        safe_text = re.sub(r'\(([^)]*\\[a-zA-Z]+[^)]*)\)', autocorrect_repl, safe_text)
        
        # --- FIX: Basic Markdown Parsing ---
        
        # Convert Markdown Images: ![alt](path) -> <img src="path">
        safe_text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1" style="max-width: 100%; height: auto; display: block; margin: 10px auto;">', safe_text)
        
        # Convert **text** to Bold
        safe_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe_text)
        # Convert *text* to Italic
        safe_text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', safe_text)
        
        # Split by double newlines for structural paragraphs
        blocks = safe_text.split('\n\n')
        processed_blocks = []
        for block in blocks:
            # Single newlines become visual line breaks ONLY outside of math
            block_content = block.replace('\n', '<br>')
            processed_blocks.append(f"<div class='block'>{block_content}</div>")
            
        final_body = "\n".join(processed_blocks)
        
        # Restore the math blocks completely untouched
        for i, math_text in enumerate(math_blocks):
            final_body = final_body.replace(f"__MATH_BLOCK_{i}__", math_text)

        # Inject Header
        header_text = self.header_input.text().strip()
        if header_text:
            safe_header = html.escape(header_text)
            final_body = f"<h2 style='text-align: center; border-bottom: 2px solid #000; padding-bottom: 10px; margin-top: 0;'>{safe_header}</h2>\n" + final_body

        # Inject Watermark
        watermark_text = getattr(self, 'watermark_input', None)
        if watermark_text and watermark_text.text().strip():
            safe_watermark = html.escape(watermark_text.text().strip())
            final_body = f"<div class='watermark'>{safe_watermark}</div>\n" + final_body

        # Build dynamic CSS
        margin = self.margin_spin.value()
        font_size = self.font_spin.value()
        line_spacing = self.line_space_spin.value()
        page_size = self.page_size_cb.currentText()

        html_template = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <script>
                MathJax = {{
                    tex: {{
                        inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
                        displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']],
                        processEscapes: true,
                        processEnvironments: true
                    }},
                    svg: {{
                        fontCache: 'global'
                    }}
                }};
            </script>
            <script id="MathJax-script" async src="{self.mathjax_url}"></script>
            <style>
                @page {{
                    size: {page_size};
                    margin: {margin}mm;
                }}
                body{{
                    font-family: 'Segoe UI', system-ui, Arial, sans-serif;
                    font-size: {font_size}pt;
                    line-height: {line_spacing};
                    color: #000;
                    background: #fff;
                    margin: 0;
                    padding: {margin}mm; /* Matches @page margin for accurate screen preview */
                    box-sizing: border-box;
                    min-height: 100vh; /* FIX: Force background to fill the preview window vertically */
                }}
                .block {{
                    margin-bottom: 1.2em;
                    page-break-inside: auto;
                }}
                .watermark{{
                    position: fixed;
                    top: 50%;
                    left: 50%;
                    transform: translate(-50%, -50%) rotate(-45deg);
                    font-size: 80pt;
                    font-weight: bold;
                    color: rgba(200, 200, 200, 0.2);
                    z-index: -1;
                    white-space: nowrap;
                    pointer-events: none;
                    user-select: none;
                }}
                /* Protect equations from page breaks */
                mjx-container[display="true"] {{
                    page-break-inside: avoid;
                   margin: 1.5em 0 !important;
                    padding: 0.5em 0;
                }}
                /* Fix white-space inheritance inside math */
                mjx-container {{
                    white-space: normal !important;
                }}
            </style>
        </head>
        <body>
            {final_body}
        </body>
        </html>
        """
        return html_template

    def update_preview(self):
        html_content = self.generate_html()
        self.web_view.setHtml(html_content, QUrl("file:///"))

    def export_pdf(self):
        default_name = self.filename_input.text().strip()
        if default_name.endswith('.pdf') or default_name.endswith('.docx'):
            default_name = os.path.splitext(default_name)[0]
        default_name += '.pdf'
        
        # Remember Last Directory
        last_dir = self.settings.value("last_dir", os.path.expanduser("~"))
        save_path = os.path.join(last_dir, default_name)
            
        filepath, _ = QFileDialog.getSaveFileName(self, "Save PDF", save_path, "PDF Files (*.pdf)")
        if not filepath:
            return
            
        self.settings.setValue("last_dir", os.path.dirname(filepath))

        # Setup precise Chromium Print Layout to match settings
        layout = QPageLayout()
        
        # Map string to QPageSize
        ps_str = self.page_size_cb.currentText()
        if ps_str == "A4":
            layout.setPageSize(QPageSize(QPageSize.A4))
        elif ps_str == "Letter":
            layout.setPageSize(QPageSize(QPageSize.Letter))
        elif ps_str == "Legal":
            layout.setPageSize(QPageSize(QPageSize.Legal))
            
        # Margins in millimeters
        m = self.margin_spin.value()
        layout.setMargins(QMarginsF(m, m, m, m))

        # Perform the export
        self.export_btn.setText("⏳ Generating...")
        self.export_btn.setEnabled(False)
        
        self.progress_bar.show()
        self.progress_bar.setFormat("Generating PDF Layout...")
        self.progress_bar.setRange(0, 0) # Infinite loop animation for Chromium background tasks
            
        self.web_view.page().printToPdf(filepath, layout)

    def _on_pdf_finished(self, saved_filepath, success):
        self.progress_bar.setFormat("%p%") # Reset format
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(100)
        QTimer.singleShot(2000, self.progress_bar.hide) # Hide it after 2 seconds
        
        self.export_btn.setText("💾 Export PDF")
        self.export_btn.setEnabled(True)
        
        if success:
            QMessageBox.information(self, "Success", f"PDF exported successfully to:\n{saved_filepath}")
            # Try to auto-open
            try:
                if os.name == 'nt': os.startfile(saved_filepath)
                elif sys.platform == 'darwin': subprocess.run(['open', saved_filepath])
                else: subprocess.run(['xdg-open', saved_filepath])
            except:
                pass
        else:
            QMessageBox.critical(self, "Error", "Failed to generate PDF. Make sure the file isn't open in another program.")

    def export_word(self):
        import docx
        from docx.shared import Pt
        
        default_name = self.filename_input.text().strip()
        if default_name.endswith('.pdf') or default_name.endswith('.docx'):
            default_name = os.path.splitext(default_name)[0]
        default_name += '.docx'
        
        # Remember Last Directory
        last_dir = self.settings.value("last_dir", os.path.expanduser("~"))
        save_path = os.path.join(last_dir, default_name)
            
        filepath, _ = QFileDialog.getSaveFileName(self, "Save Word Document", save_path, "Word Documents (*.docx)")
        if not filepath:
            return
            
        self.settings.setValue("last_dir", os.path.dirname(filepath))
            
        self.export_word_btn.setText("⏳ Generating...")
        self.export_word_btn.setEnabled(False)
        self.progress_bar.show()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        QApplication.processEvents()
        
        try:
            doc = docx.Document()
            
            # Set default font styling
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(self.font_spin.value())
            
            raw_text = self.editor.toPlainText()
            
            # --- FIX: Auto-correct AI copy-paste errors ---
            raw_text = re.sub(r'^[ \t]*\[[ \t]*$', r'\\[', raw_text, flags=re.MULTILINE)
            raw_text = re.sub(r'^[ \t]*\][ \t]*$', r'\\]', raw_text, flags=re.MULTILINE)
            
            # --- FIX: Protect Math from Word Paragraph Splitting ---
            math_blocks = []
            def math_repl(match):
                fixed_math = re.sub(r'(?<!\\)\\\s*\n', r'\\\\\n', match.group(0))
                math_blocks.append(fixed_math)
                return f"__MATH_{len(math_blocks)-1}__"

            # Extract properly formatted math FIRST
            safe_text = re.sub(r'\$\$.*?\$\$', math_repl, raw_text, flags=re.DOTALL)
            safe_text = re.sub(r'(?<!\$)\$.*?\$(?!\$)', math_repl, safe_text)
            safe_text = re.sub(r'\\\[.*?\\\]', math_repl, safe_text, flags=re.DOTALL)
            safe_text = re.sub(r'\\\(.*?\\\)', math_repl, safe_text)

            # Autocorrect stray Parentheses to LaTeX Inline Math AFTER extracting good math
            def autocorrect_repl(match):
                math_blocks.append(f"\\({match.group(1)}\\)")
                return f"__MATH_{len(math_blocks)-1}__"
                
            safe_text = re.sub(r'\(([^)]*\\[a-zA-Z]+[^)]*)\)', autocorrect_repl, safe_text)
            
            # Double newline = Structural Paragraph. Single newline = Soft Line Break.
            paragraphs = safe_text.split('\n\n')
            
            total_paras = len(paragraphs)
            for idx, para in enumerate(paragraphs):
                # Update UI Progress Bar dynamically based on paragraph count
                progress = int(((idx + 1) / max(total_paras, 1)) * 100)
                self.progress_bar.setValue(progress)
                QApplication.processEvents() # Force UI to paint the progress bar

                p = doc.add_paragraph()
                lines = para.split('\n')
                
                for line_idx, line in enumerate(lines):
                    # Parse Bold (**text**) and Italic (*text*) using Regex split
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                    
                    for part in parts:
                        if not part: continue
                        is_bold = part.startswith('**') and part.endswith('**')
                        is_italic = not is_bold and part.startswith('*') and part.endswith('*')
                        
                        clean_text = part.strip('*')
                        
                        # Restore pristine math blocks and normalize delimiters for MS Word COM
                        def restore_math(m): 
                            mtxt = math_blocks[int(m.group(1))]
                            if mtxt.startswith('\\[') and mtxt.endswith('\\]'):
                                return '$$' + mtxt[2:-2] + '$$'
                            if mtxt.startswith('\\(') and mtxt.endswith('\\)'):
                                return '$' + mtxt[2:-2] + '$'
                            return mtxt
                            
                        final_text = re.sub(r'__MATH_(\d+)__', restore_math, clean_text)
                        
                        run = p.add_run(final_text)
                        if is_bold: run.bold = True
                        if is_italic: run.italic = True
                        
                    # Add a soft carriage return if there are more lines in this structural paragraph
                    if line_idx < len(lines) - 1:
                        p.add_run().add_break()
                        
            doc.save(filepath)
            
            # --- Auto-Convert Math using Windows COM or LibreOffice Headless ---
            if getattr(self, 'auto_math_cb', None) and self.auto_math_cb.isChecked():
                word_success = False
                if os.name == 'nt':
                    self.progress_bar.setFormat("Compiling Native Word Math...")
                    QApplication.processEvents()
                    word = None
                    word_doc = None
                    try:
                        import win32com.client
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        abs_path = os.path.abspath(filepath)
                        word_doc = word.Documents.Open(abs_path)
                        
                        # Convert Math
                        for find_text, trim_len in [("$$[!$]@$$", 2), ("$[!$]@$", 1)]:
                            word.Selection.HomeKey(Unit=6)
                            word.Selection.Find.ClearFormatting()
                            while word.Selection.Find.Execute(FindText=find_text, MatchWildcards=True):
                                word.Selection.Text = word.Selection.Text[trim_len:-trim_len].strip()
                                omath = word.Selection.OMaths.Add(word.Selection.Range)
                                omath.BuildUp()
                                word.Selection.Collapse(Direction=0)
                        word_success = True
                    except Exception as e:
                        error_str = str(e)
                        if "-2147221005" in error_str or "Invalid class string" in error_str:
                            QMessageBox.warning(self, "Word Not Found", "Microsoft Word Desktop is not installed. Background math compilation failed.")
                            self.auto_math_cb.setChecked(False)
                        else:
                            print(f"Word COM Error: {e}")
                    finally:
                        if word_doc: 
                            word_doc.Save()
                            word_doc.Close()
                        if word: word.Quit()

                if not word_success:
                    # Fallback to LibreOffice Headless Conversion
                    self.progress_bar.setFormat("LibreOffice Rendering...")
                    try:
                        lo_path = "soffice"
                        if os.name == 'nt': lo_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
                        subprocess.run([lo_path, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(filepath), filepath])
                    except Exception as e:
                        print(f"LibreOffice conversion failed: {e}")
                    self.progress_bar.setFormat("LibreOffice Rendering...")
                    try:
                        import subprocess
                        # Common paths for LibreOffice
                        lo_path = "soffice" # If in PATH
                        if os.name == 'nt':
                            lo_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
                        
                        subprocess.run([lo_path, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(filepath), filepath])
                    except Exception as e:
                        print(f"LibreOffice conversion failed: {e}")
                self.progress_bar.setFormat("Compiling Native Word Math...")
                QApplication.processEvents()
                word = None
                word_doc = None
                try:
                    import win32com.client
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False # Run entirely in the background
                    abs_path = os.path.abspath(filepath)
                    word_doc = word.Documents.Open(abs_path)
                    
                    # Convert Block Math ($$ ... $$) via non-greedy wildcard
                    word.Selection.HomeKey(Unit=6)
                    word.Selection.Find.ClearFormatting()
                    while word.Selection.Find.Execute(FindText="$$[!$]@$$", MatchWildcards=True):
                        word.Selection.Text = word.Selection.Text[2:-2].strip()
                        omath = word.Selection.OMaths.Add(word.Selection.Range)
                        omath.BuildUp()
                        word.Selection.Collapse(Direction=0)

                    # Convert Inline Math ($ ... $) via non-greedy wildcard
                    word.Selection.HomeKey(Unit=6)
                    word.Selection.Find.ClearFormatting()
                    while word.Selection.Find.Execute(FindText="$[!$]@$", MatchWildcards=True):
                        word.Selection.Text = word.Selection.Text[1:-1].strip()
                        omath = word.Selection.OMaths.Add(word.Selection.Range)
                        omath.BuildUp()
                        word.Selection.Collapse(Direction=0)

                except Exception as e:
                    error_str = str(e)
                    if "-2147221005" in error_str or "Invalid class string" in error_str:
                        QMessageBox.warning(self, "Word Not Found", "Microsoft Word Desktop application is not installed or registered on this computer.\n\nThe document was saved successfully, but background math compilation requires full MS Word.")
                        if hasattr(self, 'auto_math_cb'): self.auto_math_cb.setChecked(False)
                    else:
                        QMessageBox.warning(self, "Word Automation Error", f"COM Auto-math failed:\n{e}\n\nThe file was saved, but math was not auto-converted.")
                    print(f"COM Auto-math failed: {e}")
                finally:
                    # Safely terminate the background Word process even if it crashes
                    try:
                        if word_doc:
                            word_doc.Save()
                            word_doc.Close()
                    except: pass
                    try:
                        if word:
                            word.Quit()
                    except: pass
            
            QMessageBox.information(self, "Success", f"Word Document exported successfully to:\n{filepath}")
            
            try:
                if os.name == 'nt': os.startfile(filepath)
                elif sys.platform == 'darwin': subprocess.run(['open', filepath])
                else: subprocess.run(['xdg-open', filepath])
            except:
                pass
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate Word Document:\n{e}")
            
        self.export_word_btn.setText("📝 Export Word (.docx)")
        self.export_word_btn.setEnabled(True)
        QTimer.singleShot(2000, self.progress_bar.hide)

    def export_libreoffice(self):
        import docx
        from docx.shared import Pt
        
        default_name = self.filename_input.text().strip()
        if default_name.endswith('.pdf') or default_name.endswith('.docx'):
            default_name = os.path.splitext(default_name)[0]
        default_name += '_LO.docx'
        
        last_dir = self.settings.value("last_dir", os.path.expanduser("~"))
        save_path = os.path.join(last_dir, default_name)
            
        filepath, _ = QFileDialog.getSaveFileName(self, "Save for LibreOffice", save_path, "Word Documents (*.docx)")
        if not filepath: return
            
        self.settings.setValue("last_dir", os.path.dirname(filepath))
        self.export_lo_btn.setText("⏳ Generating...")
        self.export_lo_btn.setEnabled(False)
        
        try:
            doc = docx.Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Liberation Serif' # LO preferred default
            font.size = Pt(self.font_spin.value())
            
            raw_text = self.editor.toPlainText()
            
            # --- FIX: Auto-correct AI copy-paste errors ---
            raw_text = re.sub(r'^[ \t]*\[[ \t]*$', r'\\[', raw_text, flags=re.MULTILINE)
            raw_text = re.sub(r'^[ \t]*\][ \t]*$', r'\\]', raw_text, flags=re.MULTILINE)
            
            math_blocks = []
            def math_repl(match):
                fixed_math = re.sub(r'(?<!\\)\\\s*\n', r'\\\\\n', match.group(0))
                math_blocks.append(fixed_math)
                return f"__MATH_{len(math_blocks)-1}__"

            safe_text = re.sub(r'\$\$.*?\$\$', math_repl, raw_text, flags=re.DOTALL)
            safe_text = re.sub(r'(?<!\$)\$.*?\$(?!\$)', math_repl, safe_text)
            safe_text = re.sub(r'\\\[.*?\\\]', math_repl, safe_text, flags=re.DOTALL)
            safe_text = re.sub(r'\\\(.*?\\\)', math_repl, safe_text)

            def autocorrect_repl(match):
                math_blocks.append(f"\\({match.group(1)}\\)")
                return f"__MATH_{len(math_blocks)-1}__"
                
            safe_text = re.sub(r'\(([^)]*\\[a-zA-Z]+[^)]*)\)', autocorrect_repl, safe_text)
            
            paragraphs = safe_text.split('\n\n')
            
            for para in paragraphs:
                p = doc.add_paragraph()
                lines = para.split('\n')
                
                for line_idx, line in enumerate(lines):
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                    for part in parts:
                        if not part: continue
                        is_bold = part.startswith('**') and part.endswith('**')
                        is_italic = not is_bold and part.startswith('*') and part.endswith('*')
                        clean_text = part.strip('*')
                        
                        def restore_math(m): 
                            mtxt = math_blocks[int(m.group(1))]
                            # TexMaths explicitly requires these brackets instead of $ and $$
                            if mtxt.startswith('$$') and mtxt.endswith('$$'):
                                return '\\[' + mtxt[2:-2] + '\\]'
                            if mtxt.startswith('$') and mtxt.endswith('$'):
                                return '\\(' + mtxt[1:-1] + '\\)'
                            return mtxt
                            
                        final_text = re.sub(r'__MATH_(\d+)__', restore_math, clean_text)
                        
                        run = p.add_run(final_text)
                        if is_bold: run.bold = True
                        if is_italic: run.italic = True
                        
                    if line_idx < len(lines) - 1:
                        p.add_run().add_break()
                        
            doc.save(filepath)
            
            QMessageBox.information(self, "Success", f"Document exported successfully for LibreOffice to:\n{filepath}")
            
            try:
                if os.name == 'nt': os.startfile(filepath)
                elif sys.platform == 'darwin': subprocess.run(['open', filepath])
                else: subprocess.run(['xdg-open', filepath])
            except: pass
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate LO Document:\n{e}")
            
        self.export_lo_btn.setText("📝 Export LibreOffice")
        self.export_lo_btn.setEnabled(True)

    def scan_image_ocr(self):
        import threading
        
        # Remember Last Directory
        last_dir = self.settings.value("last_dir", os.path.expanduser("~"))
        filepath, _ = QFileDialog.getOpenFileName(self, "Select Document", last_dir, "Documents & Images (*.pdf *.png *.jpg *.jpeg)")
        if not filepath:
            return
            
        self.settings.setValue("last_dir", os.path.dirname(filepath))

        # Set UI to Loading safely
        loading_msg = "\n\n[... Local OCR scanning... Note: Potato-PC OCR extracts raw text but cannot magically build LaTeX matrices ...]\n"
        self.editor.insertPlainText(loading_msg)
        
        self.progress_bar.show()
        self.progress_bar.setRange(0, 100) 
        self.progress_bar.setValue(0)
        self.progress_bar.setFormat("Loading AI Engine... 0%")
        
        # Setup continuous update timer for single images
        self.fake_progress = 0
        if hasattr(self, 'fake_timer'):
            self.fake_timer.stop()
            self.fake_timer.deleteLater()
        self.fake_timer = QTimer(self)
        
        def tick_progress():
            if self.fake_progress < 95:
                increment = max(1, (96 - self.fake_progress) // 10)
                self.fake_progress += increment
                self.progress_bar.setValue(self.fake_progress)
                self.progress_bar.setFormat(f"Scanning Image... {self.fake_progress}%")
                
        self.fake_timer.timeout.connect(tick_progress)
        
        def run_ocr():
            try:
                import fitz  # PyMuPDF
                
                # --- SPEED OPTIMIZATION ---
                if self.ocr_engine is None:
                    from rapidocr_onnxruntime import RapidOCR
                    self.ocr_engine = RapidOCR()
                    
                engine = self.ocr_engine
                extracted_text = ""
                
                if filepath.lower().endswith('.pdf'):
                    doc = fitz.open(filepath)
                    total_pages = len(doc)
                    
                    QTimer.singleShot(0, lambda: self.progress_bar.setFormat("Scanning PDF: %p%"))
                    QTimer.singleShot(0, lambda: self.progress_bar.setValue(0))
                    
                    for page_num in range(total_pages):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(dpi=150) # Extract page as image
                        img_data = pix.tobytes("png")
                        
                        # Continuous update: tick forward halfway when image frame is extracted
                        progress_half = int(((page_num + 0.5) / total_pages) * 100)
                        QTimer.singleShot(0, lambda p=progress_half: self.progress_bar.setValue(p))
                        
                        result, _ = engine(img_data)
                        if result:
                            extracted_text += f"\n--- Page {page_num+1} ---\n"
                            for line in result:
                                extracted_text += line[1] + "\n"
                                
                        # Update full progress after page OCR is finished
                        progress = int(((page_num + 1) / total_pages) * 100)
                        QTimer.singleShot(0, lambda p=progress: self.progress_bar.setValue(p))
                        
                    doc.close()
                else:
                    # Single image takes a few seconds, use continuous simulated timer
                    QTimer.singleShot(0, lambda: self.fake_timer.start(100))
                    result, _ = engine(filepath)
                    QTimer.singleShot(0, self.fake_timer.stop)
                    
                    if result:
                        for line in result:
                            extracted_text += line[1] + "\n"
                
                if not extracted_text:
                    extracted_text = "[No readable text found]"
                
                # Safely update UI from background thread
                def on_success():
                    self.fake_timer.stop()
                    text = self.editor.toPlainText()
                    text = text.replace(loading_msg, "")
                    self.editor.setPlainText(text)
                    self.editor.insertPlainText("\n" + extracted_text.strip() + "\n")
                    self.progress_bar.setFormat("Done! 100%")
                    self.progress_bar.setValue(100)
                    QTimer.singleShot(2000, self.progress_bar.hide)
                QTimer.singleShot(0, on_success)
                
            except Exception as e:
                def on_error():
                    self.fake_timer.stop()
                    text = self.editor.toPlainText()
                    text = text.replace(loading_msg, f"\n[Local OCR Error: {str(e)}]\n")
                    self.editor.setPlainText(text)
                    self.progress_bar.setFormat("Error!")
                    self.progress_bar.hide()
                QTimer.singleShot(0, on_error)

        threading.Thread(target=run_ocr, daemon=True).start()

    def scan_image_gemini(self):
        import threading
        from PIL import Image

        api_key = self.api_key_input.text().strip()
        if not api_key:
            QMessageBox.warning(self, "API Key Missing", "Please enter your Gemini API Key first.")
            return

        last_dir = self.settings.value("last_dir", os.path.expanduser("~"))
        filepath, _ = QFileDialog.getOpenFileName(self, "Select Document", last_dir, "Images (*.png *.jpg *.jpeg)")
        if not filepath: return
        self.settings.setValue("last_dir", os.path.dirname(filepath))

        loading_msg = "\n\n[... Asking Gemini to extract and format LaTeX Math...]\n"
        self.editor.insertPlainText(loading_msg)
        self.progress_bar.show()
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setFormat("Analyzing with Gemini...")

        def run_gemini():
            try:
                from google import genai
                client = genai.Client(api_key=api_key)
                img = Image.open(filepath)
                
                # A highly specific prompt to force correct LaTeX delimiting
                prompt = "Extract all text and mathematical formulas from this image. Output pure Markdown. Wrap inline math with single $ and block math/matrices with double $$. Do NOT use markdown codeblock wrappers like ```markdown around the whole response."
                
                response = client.models.generate_content(
                    model='gemini-1.5-flash',
                    contents=[prompt, img]
                )
                
                def on_success():
                    text = self.editor.toPlainText().replace(loading_msg, "")
                    self.editor.setPlainText(text)
                    self.editor.insertPlainText("\n" + response.text.strip() + "\n")
                    self.progress_bar.setFormat("%p%")
                    self.progress_bar.setRange(0, 100)
                    self.progress_bar.setValue(100)
                    QTimer.singleShot(2000, self.progress_bar.hide)
                QTimer.singleShot(0, on_success)
            except Exception as e:
                def on_error():
                    text = self.editor.toPlainText().replace(loading_msg, f"\n[Gemini Error: {str(e)}]\n")
                    self.editor.setPlainText(text)
                    self.progress_bar.hide()
                QTimer.singleShot(0, on_error)

        threading.Thread(target=run_gemini, daemon=True).start()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = MathPdfMaker()
    window.show()
    sys.exit(app.exec())
